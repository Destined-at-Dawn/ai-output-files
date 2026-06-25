# -*- coding: utf-8 -*-
"""
生成微信聊天记录读取工具包 v2.0 使用教程（docx 格式）
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

OUTPUT = r"E:\ai产出文件\牛马\mutual\mutual\projects\proj-1779089173658-j5tg2m\outputs\微信聊天记录读取工具包-使用教程.docx"

doc = Document()

# ===== 样式设置 =====
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.35
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = '微软雅黑'
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if level == 1:
        h.font.size = Pt(22)
    elif level == 2:
        h.font.size = Pt(16)
    else:
        h.font.size = Pt(13)

# ===== 辅助函数 =====
def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    # 灰色背景用段落底纹
    from docx.oxml import OxmlElement
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    p.paragraph_format.element.get_or_add_pPr().append(shd)
    return p

def add_note(text, bold_prefix=''):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()  # 表后空行

# ===== 封面 =====
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('微信聊天记录读取工具包')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('v2.0  使用教程')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('从已解密的微信数据库中提取聊天记录、联系人、\n转账红包记录，并进行文风分析')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026-06-16')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ===== 目录 =====
doc.add_heading('目录', level=1)
toc_items = [
    '1. 这个工具包能做什么',
    '2. 环境准备',
    '3. 获取解密后的数据库（前置步骤）',
    '4. 工具一：全量数据提取',
    '5. 工具二：导出特定聊天记录',
    '6. 工具三：生成分析报告',
    '7. 工具四：文风DNA分析',
    '8. 工具五/六：提取红包和转账',
    '9. 微信数据库结构速查',
    '10. 常见问题 FAQ',
    '11. 免责声明',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
doc.add_page_break()

# ===== 第1章 =====
doc.add_heading('1. 这个工具包能做什么', level=1)

doc.add_paragraph(
    '微信聊天记录读取工具包是一套 Python 脚本，用于从已解密的微信数据库中提取和分析数据。'
    '它包含 6 个工具，覆盖了从原始数据库到结构化分析的完整流程。'
)

add_table(
    ['脚本', '功能', '输出'],
    [
        ['extract_full_data.py', '全量数据提取（聊天+联系人+群成员）', 'corpus.json, all_sessions.json 等'],
        ['export_chat.py', '单个联系人/群聊导出为 Markdown', 'chat_xxx.md'],
        ['analyze_wechat.py', '聊天分类汇总 + 统计报告', 'MD + JSON 报告'],
        ['analyze_voice_dna.py', '文风DNA蒸馏（私聊vs群聊风格对比）', '文风分析文本'],
        ['extract_red_envelopes.py', '红包记录提取（含金额）', 'red_envelopes.json'],
        ['extract_transfer_amounts.py', '转账记录提取（含金额）', 'transfers.json'],
    ]
)

add_note('这些工具只处理已解密的数据库，不涉及任何解密操作。', '注意：')

# ===== 第2章 =====
doc.add_heading('2. 环境准备', level=1)

doc.add_heading('2.1 安装 Python', level=2)
doc.add_paragraph('你需要 Python 3.8 或更高版本（推荐 3.10+）。')
doc.add_paragraph('下载地址：https://www.python.org/downloads/')
add_note('安装时勾选 "Add Python to PATH"。', '重要：')

doc.add_heading('2.2 安装依赖', level=2)
doc.add_paragraph('打开命令行（CMD 或 PowerShell），执行：')
add_code('pip install zstandard')
add_note('zstandard 用于解压微信 4.0 的 zstd 压缩消息。如果安装失败，试试：')
add_code('pip install zstandard --only-binary=:all:')

doc.add_heading('2.3 下载工具包', level=2)
doc.add_paragraph('解压 wechat-reader-toolkit-v2.0.zip 到任意目录，你会看到：')
add_code(
    'wechat-reader-toolkit/\n'
    '  extract_full_data.py      # 全量提取\n'
    '  export_chat.py            # 聊天导出\n'
    '  analyze_wechat.py         # 分析报告\n'
    '  analyze_voice_dna.py      # 文风分析\n'
    '  extract_red_envelopes.py  # 红包提取\n'
    '  extract_transfer_amounts.py  # 转账提取\n'
    '  README.md                 # 英文说明'
)

# ===== 第3章 =====
doc.add_heading('3. 获取解密后的数据库（前置步骤）', level=1)

doc.add_paragraph(
    '本工具包只读取已解密的 SQLite 数据库。你需要先使用第三方工具解密微信数据库。'
)

doc.add_heading('3.1 推荐工具', level=2)

add_table(
    ['工具', '安装方式', '适用场景'],
    [
        ['PyWxDump', 'pip install pywxdump', '最主流，支持 Windows/macOS，有 GUI'],
        ['wx_key', 'pip install wx_key', '轻量，需微信正在运行，通过 DLL Hook 提取密钥'],
    ]
)

doc.add_heading('3.2 解密流程（以 PyWxDump 为例）', level=2)

add_code(
    '# 1. 安装\n'
    'pip install pywxdump\n\n'
    '# 2. 获取数据库加密密钥（需要微信正在运行）\n'
    'pywxdump bias_addr\n\n'
    '# 3. 解密数据库\n'
    'pywxdump decrypt -k <上一步获取的密钥> -i <加密数据库目录> -o <输出目录>\n\n'
    '# 4. 解密完成后，输出目录中会有：\n'
    '#    message_0.db    -- 聊天记录（几百MB）\n'
    '#    contact.db      -- 联系人信息\n'
    '#    general.db      -- 转账/红包/收藏\n'
    '#    session.db      -- 会话列表'
)

doc.add_heading('3.3 数据库文件说明', level=2)

add_table(
    ['文件', '大小', '内容'],
    [
        ['message_0.db', '几百MB', '所有聊天消息（按 hash 分表）'],
        ['contact.db', '几MB', '联系人/群成员信息'],
        ['general.db', '几十MB', '转账/红包/收藏/小程序等'],
        ['session.db', '几MB', '会话列表（最近聊天）'],
    ]
)

add_note('四个文件必须在同一个目录下。', '重要：')

# ===== 第4章 =====
doc.add_heading('4. 工具一：全量数据提取', level=1)

doc.add_paragraph(
    '这是最常用的工具。它会遍历数据库中的所有消息，提取聊天记录、联系人、群成员、转账红包等数据，'
    '输出一组 JSON 文件和一份 Markdown 报告。'
)

doc.add_heading('4.1 用法', level=2)
add_code(
    'python extract_full_data.py --db-dir "你的数据库目录" --output "输出目录"'
)

doc.add_heading('4.2 参数说明', level=2)
add_table(
    ['参数', '必填', '说明'],
    [
        ['--db-dir', '是', '解密后的数据库目录（含 message_0.db, contact.db 等）'],
        ['--output', '是', '输出目录（会自动创建）'],
        ['--wxid', '否', '你的微信 wxid（用于识别"自己"发的消息）'],
    ]
)

doc.add_heading('4.3 输出文件', level=2)
add_table(
    ['文件', '内容', '格式'],
    [
        ['all_sessions.json', '所有会话统计（按消息数排序）', 'JSON 数组'],
        ['top_private.json', '私聊排行 Top 20', 'JSON 数组'],
        ['top_groups.json', '群聊排行 Top 20', 'JSON 数组'],
        ['corpus.json', '全量消息语料（33000+ 条）', 'JSON 数组'],
        ['transfers.json', '转账记录', 'JSON 对象'],
        ['red_envelopes.json', '红包记录', 'JSON 对象'],
        ['group_members.json', '群成员列表', 'JSON 对象'],
        ['data_summary.md', 'Markdown 分析报告', 'Markdown'],
    ]
)

doc.add_heading('4.4 实际例子', level=2)
add_code(
    '# 假设解密后的数据库在 D:\\wechat-decrypted\\ 目录下\n'
    '# 输出到当前目录的 output 文件夹\n\n'
    'python extract_full_data.py \\\n'
    '    --db-dir "D:\\wechat-decrypted" \\\n'
    '    --output "output" \\\n'
    '    --wxid "wxid_abc123"\n\n'
    '# 完成后 output 目录下会有 8 个文件\n'
    '# data_summary.md 可以直接用 Markdown 阅读器打开'
)

# ===== 第5章 =====
doc.add_heading('5. 工具二：导出特定聊天记录', level=1)

doc.add_paragraph(
    '如果你只想导出某个联系人或群聊的聊天记录为 Markdown 文件，用这个工具。'
)

doc.add_heading('5.1 用法', level=2)
add_code(
    'python export_chat.py \\\n'
    '    --db "message_0.db的路径" \\\n'
    '    --contact-db "contact.db的路径" \\\n'
    '    --target "对方的wxid或群聊ID" \\\n'
    '    --output "输出文件.md"'
)

doc.add_heading('5.2 参数说明', level=2)
add_table(
    ['参数', '必填', '说明'],
    [
        ['--db', '是', 'message_0.db 的完整路径'],
        ['--contact-db', '否', 'contact.db 路径（用于解析联系人名称）'],
        ['--target', '是', '导出目标：联系人 wxid 或群聊 ID（xxx@chatroom）'],
        ['--target-name', '否', '目标显示名（默认用 wxid）'],
        ['--output', '是', '输出 Markdown 文件路径'],
        ['--sender-map', '否', '发送者映射 JSON（格式: {"2":"我","对方wxid":"名字"}）'],
    ]
)

doc.add_heading('5.3 如何找到 wxid 和群聊 ID？', level=2)
doc.add_paragraph(
    '方法一：先运行 extract_full_data.py，查看 all_sessions.json，里面列出了所有联系人/群聊的 ID。\n\n'
    '方法二：查看 session.db 的 session 表，username 列就是 ID。'
)

doc.add_heading('5.4 实际例子', level=2)
add_code(
    '# 导出和"张三"的私聊记录\n'
    'python export_chat.py \\\n'
    '    --db "D:\\wechat-decrypted\\message_0.db" \\\n'
    '    --contact-db "D:\\wechat-decrypted\\contact.db" \\\n'
    '    --target "wxid_zhangsan123" \\\n'
    '    --target-name "张三" \\\n'
    '    --output "chat_张三.md"\n\n'
    '# 导出群聊记录\n'
    'python export_chat.py \\\n'
    '    --db "D:\\wechat-decrypted\\message_0.db" \\\n'
    '    --target "12345678@chatroom" \\\n'
    '    --target-name "我的群聊" \\\n'
    '    --output "chat_我的群聊.md"'
)

# ===== 第6章 =====
doc.add_heading('6. 工具三：生成分析报告', level=1)

doc.add_paragraph(
    '这个工具会自动将聊天消息按类别分类（学习/AI/竞赛/生活等），生成结构化的分析报告。'
)

doc.add_heading('6.1 用法', level=2)
add_code(
    'python analyze_wechat.py --db-dir "数据库目录" --out-dir "输出目录"'
)

doc.add_heading('6.2 参数说明', level=2)
add_table(
    ['参数', '必填', '说明'],
    [
        ['--db-dir', '是', '数据库目录（含 message_0.db 等）'],
        ['--out-dir', '是', '输出目录'],
        ['--top-n', '否', '分析前 N 个联系人/群（默认 10）'],
        ['--verbose', '否', '输出详细分类过程日志'],
    ]
)

doc.add_heading('6.3 输出', level=2)
doc.add_paragraph(
    '生成 Markdown 报告文件，包含：\n'
    '- 联系人/群聊的消息统计排行\n'
    '- 每个会话的消息分类汇总（按话题归类）\n'
    '- 可直接用浏览器或 Markdown 编辑器查看'
)

# ===== 第7章 =====
doc.add_heading('7. 工具四：文风DNA分析', level=1)

doc.add_paragraph(
    '这个工具分析你的私聊和群聊用语差异，生成"文风DNA报告"。'
    '包括消息长度分布、标点使用、高频词等维度。'
)

doc.add_heading('7.1 用法', level=2)
add_code(
    'python analyze_voice_dna.py --corpus "corpus.json的路径" --output "文风报告.txt"'
)

add_note('需要先运行 extract_full_data.py 生成 corpus.json。', '前置条件：')

doc.add_heading('7.2 参数说明', level=2)
add_table(
    ['参数', '必填', '说明'],
    [
        ['--corpus', '是', 'corpus.json 文件路径'],
        ['--sessions', '否', 'top_private.json 路径（默认从 corpus 同目录查找）'],
        ['--output', '否', '输出文件路径（默认打印到终端）'],
    ]
)

# ===== 第8章 =====
doc.add_heading('8. 工具五/六：提取红包和转账', level=1)

doc.add_paragraph(
    '这两个工具分别从 general.db 中提取红包和转账记录，包括金额信息。'
)

doc.add_heading('8.1 红包提取', level=2)
add_code(
    'python extract_red_envelopes.py --db-dir "数据库目录" --output "red_envelopes.json"'
)

doc.add_heading('8.2 转账提取', level=2)
add_code(
    'python extract_transfer_amounts.py --db-dir "数据库目录" --output "transfers.json"'
)

doc.add_heading('8.3 输出格式', level=2)
doc.add_paragraph(
    '两个工具都输出 JSON 文件，包含：\n'
    '- 收发双方（wxid + 昵称/备注）\n'
    '- 金额（如有）\n'
    '- 时间戳\n'
    '- 货币类型'
)

# ===== 第9章 =====
doc.add_heading('9. 微信数据库结构速查', level=1)

doc.add_paragraph('以下是微信 4.0 解密后的数据库关键结构，供高级用户参考。')

doc.add_heading('9.1 message_0.db 消息表', level=2)
add_table(
    ['列名', '类型', '说明'],
    [
        ['local_id', 'INTEGER', '本地消息 ID'],
        ['create_time', 'INTEGER', 'Unix 时间戳（秒）'],
        ['real_sender_id', 'INTEGER', '发送者 ID（2 = 数据库主人）'],
        ['message_content', 'TEXT/BLOB', '消息内容（可能是 zstd 压缩）'],
        ['compress_content', 'BLOB', 'zstd 压缩的内容（magic: 0x28B52FFD）'],
        ['status', 'INTEGER', '消息状态'],
        ['local_type', 'INTEGER', '10000 = 系统消息'],
    ]
)

add_note(
    '消息按 Msg_md5_hash 分表，表名格式为 Msg_00 ~ Msg_xx。'
    '需要遍历所有表名来查询。',
    '注意：'
)

doc.add_heading('9.2 contact.db 联系人表', level=2)
add_table(
    ['表名', '关键列', '说明'],
    [
        ['contact', 'username, nick_name, remark, alias', '联系人基本信息'],
        ['chatroom_member', 'room_id, member_id', '群聊成员关系'],
        ['chat_room', 'id, username, owner', '群聊基本信息'],
    ]
)

doc.add_heading('9.3 general.db 转账/红包', level=2)
add_table(
    ['表名', '说明'],
    [
        ['transferTable', '转账记录'],
        ['redEnvelopeTable', '红包记录'],
    ]
)

doc.add_heading('9.4 session.db 会话列表', level=2)
doc.add_paragraph('session 表包含 username（联系人/群聊 ID）和最新消息时间。')

# ===== 第10章 =====
doc.add_heading('10. 常见问题 FAQ', level=1)

faqs = [
    ('Q: zstandard 安装失败怎么办？',
     'A: 尝试只安装预编译版本：pip install zstandard --only-binary=:all:\n'
     '如果仍然失败，确认 Python 版本 >= 3.8。'),
    ('Q: 数据库是加密的，打不开怎么办？',
     'A: 本工具只处理已解密的数据库。请先用 PyWxDump 或 wx_key 解密。'
     '详见第 3 章。'),
    ('Q: 不知道数据库文件在哪里？',
     'A: 微信 4.0 默认路径通常在：\n'
     '  D:\\data\\wechat\\xwechat_files\\wxid_xxx\\db_storage\\message\\message_0.db\n'
     '或：%LOCALAPPDATA%\\Packages\\...\\LocalCache\\Roaming\\Tencent\\WeChat\\...\n'
     '也可以用 Everything 搜索 "message_0.db"。'),
    ('Q: 消息内容显示为乱码 / bytes 怎么办？',
     'A: 微信 4.0 的 message_content 可能是 zstd 压缩的。'
     '脚本已内置解压逻辑，确保安装了 zstandard 包即可。'),
    ('Q: 群聊消息只有 wxid 没有名字？',
     'A: 需要同时提供 contact.db（通过 --contact-db 参数），'
     '脚本会从 chatroom_member 表中解析群成员名称。'),
    ('Q: 找不到消息分表（Msg_xx）？',
     'A: 微信消息按 MD5 分表，表名格式为 Msg_00 ~ Msg_xx。'
     '不同数据库的分表数量不同。脚本会自动遍历查询。'),
    ('Q: 运行时报 "no such table" 错误？',
     'A: 确认你传入的是正确的数据库文件。\n'
     '  - 聊天消息 → message_0.db\n'
     '  - 联系人 → contact.db\n'
     '  - 转账/红包 → general.db\n'
     '不要混淆这四个文件。'),
    ('Q: 导出的聊天记录不完整？',
     'A: 可能原因：\n'
     '  1. 数据库解密不完整（重新解密）\n'
     '  2. 对方的 wxid 输入有误（检查 all_sessions.json）\n'
     '  3. 群聊消息被分到了不同的分表（脚本会自动处理）'),
]

for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    run.font.size = Pt(11)
    p = doc.add_paragraph(a)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)

# ===== 第11章 =====
doc.add_heading('11. 免责声明', level=1)

doc.add_paragraph(
    '本工具仅供个人数据分析使用。使用本工具时，你应当：\n\n'
    '1. 遵守当地法律法规\n'
    '2. 仅分析自己的数据，或获得数据所有者明确授权的数据\n'
    '3. 尊重他人隐私，不将分析结果用于非法用途\n'
    '4. 理解数据库解密可能涉及法律风险，请自行评估\n\n'
    '作者不对因使用本工具产生的任何后果承担责任。'
)

# ===== 保存 =====
doc.save(OUTPUT)
print(f'[OK] Tutorial saved: {OUTPUT}')
print(f'     Size: {__import__("os").path.getsize(OUTPUT)/1024:.1f} KB')
