#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
道法术器Word文档生成器
功能：将Markdown格式的道法术器内容转换为Word文档
"""

import sys
import io
import re

# 修复Windows控制台编码问题
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')


def parse_markdown(md_content: str) -> dict:
    """解析Markdown内容"""
    result = {
        'title': '',
        'sections': []
    }

    lines = md_content.split('\n')
    current_section = None
    current_items = []

    for line in lines:
        line = line.strip()

        # 解析标题
        if line.startswith('# ') and not result['title']:
            result['title'] = line[2:].strip()
            continue

        # 识别新章节
        if line.startswith('## '):
            # 保存上一个章节
            if current_section and current_items:
                result['sections'].append({
                    'title': current_section,
                    'items': current_items
                })

            current_section = line[3:].strip()
            current_items = []
        elif line.startswith('### '):
            # 保存上一项
            if current_items:
                result['sections'].append({
                    'title': current_section,
                    'items': current_items
                })

            current_section = line[4:].strip()
            current_items = []
        elif current_section:
            if line.startswith('-') or line.startswith('*'):
                content = line[1:].strip()
                if content.startswith('**'):
                    # 标题项
                    if current_items:
                        result['sections'].append({
                            'title': current_section,
                            'items': current_items
                        })
                    current_section = content.replace('**', '').strip()
                    current_items = []
                else:
                    current_items.append(content)
            elif line.startswith('|'):
                # 表格行
                current_items.append(line)
            elif line and not line.startswith('>') and not line.startswith('---'):
                current_items.append(line)

    # 保存最后一个章节
    if current_section and current_items:
        result['sections'].append({
            'title': current_section,
            'items': current_items
        })

    return result


def generate_docx(parsed_data: dict, output_file: str):
    """生成Word文档"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("⚠️  需要安装python-docx库")
        print("   请运行: pip install python-docx")
        return False

    doc = Document()

    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)

    # 标题
    title = doc.add_heading(parsed_data['title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 遍历章节
    for section in parsed_data['sections']:
        # 添加章节标题
        doc.add_heading(section['title'], level=2)

        # 添加内容
        for item in section['items']:
            if item.startswith('|'):
                # 表格
                continue
            elif item.startswith('['):
                # 复选框
                p = doc.add_paragraph(item, style='List Bullet')
            else:
                doc.add_paragraph(item)

        doc.add_paragraph()  # 空行

    # 保存文档
    doc.save(output_file)
    return True


def main():
    if len(sys.argv) < 2:
        print("用法: python generate_docx.py <markdown文件路径>")
        print("示例: python generate_docx.py transcript_cleaned.md")
        sys.exit(1)

    md_file = sys.argv[1]

    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # 解析Markdown
        parsed_data = parse_markdown(md_content)

        # 生成输出文件名
        output_file = md_file.replace('.md', '.docx').replace('_cleaned', '')

        # 生成Word文档
        if generate_docx(parsed_data, output_file):
            print(f"✅ Word文档生成完成！")
            print(f"📁 输出文件: {output_file}")
        else:
            print("❌ Word文档生成失败")

    except Exception as e:
        print(f"❌ 错误: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
