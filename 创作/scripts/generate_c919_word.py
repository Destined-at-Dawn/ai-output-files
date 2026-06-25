import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


PROJECT_NAME = "光影铸翼：C919的飞天皮影志"
PROJECT_ID = "XLY00024"

CONSISTENCY_REPLACEMENTS = {
    "老教授": "老教师",
    "招飞面试官": "招飞人员",
    "面试官": "招飞人员",
    "记者、儿童、老航空人": "女记者、儿童、老航天员",
    "记者、儿童、老航天员": "女记者、儿童、老航天员",
    "老航空人": "老航天员",
    "记者记录时代": "女记者记录时代",
    "记者的话筒": "女记者的话筒",
}


def normalize_text(text: str) -> str:
    for old, new in CONSISTENCY_REPLACEMENTS.items():
        text = text.replace(old, new)

    text = text.replace("人物统一采用至少 6 个关节设计", "人物统一采用 6 个关节设计")
    text = text.replace("人物至少保留 6 个关节", "所有人物均有且仅有 6 个关节")
    text = text.replace("至少 6 个关节", "6 个关节")
    text = text.replace("人物至少6个关节", "人物有且仅有6个关节")
    text = text.replace("关节为颈、双肩、双肘、腰。", "关节为 6 个：颈部、左肩、右肩、左肘、右肘、腰部。")

    if text.startswith("本剧每一幕都设计了 3 个相关人物角色。"):
        text = (
            "本剧每一幕都设计了 3 个相关人物角色。所有人物角色均有且仅有 6 个关节："
            "颈部、左肩、右肩、左肘、右肘、腰部。腿部不设置独立活动关节，"
            "以降低展示版本的制作风险，并保证后续舵机控制稳定性。"
        )
    if text.startswith("所有人物均有且仅有 6 个关节："):
        text = "所有人物均有且仅有 6 个关节：颈部、左肩、右肩、左肘、右肘、腰部。"

    return text


def repair_paragraphs(paragraphs: list[str]) -> list[str]:
    repaired: list[str] = []
    i = 0
    while i < len(paragraphs):
        text = paragraphs[i]

        if text.startswith("1. 青年学生 "):
            repaired.append(text)
            if i + 1 < len(paragraphs) and paragraphs[i + 1].startswith("抬头凝望、合上书本。关节为"):
                i += 2
                continue

        if text.startswith("5. 总设计师 "):
            repaired.append(
                "5. 总设计师 出现于第二幕。造型采用白发、图纸、工程外套等元素，"
                "代表总体设计与安全把关。动作重点是看图纸、标注、抬手指向关键结构。"
                "关节为 6 个：颈部、左肩、右肩、左肘、右肘、腰部。"
            )
            if i + 1 < len(paragraphs) and paragraphs[i + 1].startswith("全把关。"):
                i += 2
                continue

        if text.startswith("6. 系统协调工程师 "):
            repaired.append(
                "6. 系统协调工程师 出现于第二幕。造型配电话、记录板，"
                "代表多部门联调中的沟通枢纽。动作重点是接打电话、低头记录参数、"
                "抬手示意确认。关节为 6 个：颈部、左肩、右肩、左肘、右肘、腰部。"
            )
            if i + 1 < len(paragraphs) and paragraphs[i + 1].startswith("。动作重点是"):
                i += 2
                continue

        if not (
            text.startswith("1. 青年学生 ")
            or text.startswith("5. 总设计师 ")
            or text.startswith("6. 系统协调工程师 ")
        ):
            repaired.append(text)
        i += 1
    return repaired


def clear_cell(cell) -> None:
    tc = cell._tc
    tc_pr = tc.tcPr
    for child in list(tc):
        if child is not tc_pr:
            tc.remove(child)
    if tc_pr is None:
        tc.insert(0, OxmlElement("w:tcPr"))


def replace_paragraph_text(paragraph, text: str, bold: bool = False) -> None:
    p = paragraph._p
    p_pr = p.pPr
    for child in list(p):
        if child is not p_pr:
            p.remove(child)
    run = paragraph.add_run(text)
    run.bold = bold


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    replace_paragraph_text(paragraph, text)


def set_run_font(run, size=12, font_name="仿宋") -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def format_paragraph(paragraph, is_heading: bool = False) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = None if is_heading else Pt(24)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5


HEADINGS = {
    "1. 主题介绍",
    "2. 剧情策划",
    "2.1 故事主线",
    "2.2 故事情节",
    "2.3 剧目场次",
    "2.4 角色设计",
    "2.4.1 角色形象",
    "2.4.2 角色动作设计",
    "2.4.3 角色加工",
    "2.5 皮影表演策划",
    "加工要求：",
    "第一幕动作组：",
    "第二幕动作组：",
    "第三幕动作组：",
    "第四幕动作组：",
    "第一幕：《梦想启航》",
    "第二幕：《灯火攻关》",
    "第三幕：《首飞时刻》",
    "第四幕：《盛世见证》",
    "场景一：航空课堂",
    "场景二：招飞初选现场",
    "场景：设计室与装配工位",
    "场景：C919 驾驶舱",
    "场景：开阔天空与公共观看空间",
}


def is_heading(text: str) -> bool:
    if text in HEADINGS:
        return True
    if text[:3].isdigit() and "." in text[:5]:
        return True
    return False


def normalize_structure(paragraphs: list[str]) -> list[str]:
    normalized = []
    for text in paragraphs:
        if text == "主题介绍":
            normalized.append("1. 主题介绍")
        elif text == "剧情策划":
            normalized.append("2. 剧情策划")
        else:
            normalized.append(text)
    return normalized


def main() -> None:
    template_path = Path(sys.argv[1])
    source_path = Path(sys.argv[2])
    output_path = Path(sys.argv[3])

    source = Document(str(source_path))
    source_cell = source.tables[0].rows[1].cells[0]
    body_paragraphs = normalize_structure(repair_paragraphs([normalize_text(p.text.strip()) for p in source_cell.paragraphs]))

    doc = Document(str(template_path))
    table = doc.tables[0]
    set_cell_text(table.rows[0].cells[1], PROJECT_NAME)
    set_cell_text(table.rows[0].cells[3], PROJECT_ID)

    body_cell = table.rows[1].cells[0]
    template_paragraph = body_cell.paragraphs[0]
    clear_cell(body_cell)
    for text in body_paragraphs:
        p = body_cell.add_paragraph(style=template_paragraph.style)
        if text:
            heading = is_heading(text)
            replace_paragraph_text(p, text, bold=heading)
            format_paragraph(p, is_heading=heading)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.space_after = Pt(6)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(output_path)


if __name__ == "__main__":
    main()
