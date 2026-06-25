import shutil
import sys
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document


PROJECT_NAME = "光影铸翼：C919的飞天皮影志"
PROJECT_ID = "XLY00024"


def clear_paragraph_runs(paragraph) -> None:
    p = paragraph._p
    p_pr = p.pPr
    for child in list(p):
        if child is not p_pr:
            p.remove(child)


def fill_paragraph(paragraph, text: str, template_run=None) -> None:
    clear_paragraph_runs(paragraph)
    if not text:
        return
    run = paragraph.add_run(text)
    if template_run is not None and template_run._r.rPr is not None:
        run._r.insert(0, deepcopy(template_run._r.rPr))


def clear_cell_keep_properties(cell) -> None:
    tc = cell._tc
    tc_pr = tc.tcPr
    for child in list(tc):
        if child is not tc_pr:
            tc.remove(child)


def append_paragraph_like(cell, template_paragraph, text: str) -> None:
    new_p = deepcopy(template_paragraph._p)
    cell._tc.append(new_p)
    paragraph = cell.paragraphs[-1]
    run = template_paragraph.runs[0] if template_paragraph.runs else None
    fill_paragraph(paragraph, text, template_run=run)


def main() -> None:
    template_path = Path(sys.argv[1])
    source_path = Path(sys.argv[2])
    output_path = Path(sys.argv[3])

    backup_path = output_path.with_name(
        f"{output_path.stem}_before_template_fill_{datetime.now().strftime('%Y%m%d_%H%M%S')}{output_path.suffix}"
    )
    if output_path.exists():
        shutil.copy2(output_path, backup_path)

    template = Document(str(template_path))
    source = Document(str(source_path))

    source_cell = source.tables[0].rows[1].cells[0]
    body_texts = [p.text for p in source_cell.paragraphs]

    table = template.tables[0]
    name_para = table.rows[0].cells[1].paragraphs[0]
    id_para = table.rows[0].cells[3].paragraphs[0]
    fill_paragraph(name_para, PROJECT_NAME, template_run=name_para.runs[0] if name_para.runs else None)
    fill_paragraph(id_para, PROJECT_ID, template_run=id_para.runs[0] if id_para.runs else None)

    body_cell = table.rows[1].cells[0]
    template_paragraphs = list(body_cell.paragraphs)
    template_for_extra = template_paragraphs[-1]
    clear_cell_keep_properties(body_cell)
    for idx, text in enumerate(body_texts):
        model = template_paragraphs[idx] if idx < len(template_paragraphs) else template_for_extra
        append_paragraph_like(body_cell, model, text)

    template.save(str(output_path))
    print(f"UPDATED={output_path}")
    if backup_path.exists():
        print(f"BACKUP={backup_path}")


if __name__ == "__main__":
    main()
