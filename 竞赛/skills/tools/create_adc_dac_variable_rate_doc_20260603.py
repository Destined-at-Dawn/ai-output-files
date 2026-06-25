from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PKG = Path(r"D:\AMD\老板交付_ADC1采样DAC1可调采样率回放_默认4kSPS_20260602")
OUT_NAME = "实验指导书_ADC1采样DAC1可调采样率回放_默认4kSPS.docx"
OUT_DOC = PKG / OUT_NAME
PROJECT_OUTPUT = Path(r"D:\AMD\project_output\实验文档") / OUT_NAME

ADC_IMG = PKG / "references" / "0178A2918E7353AF817F6876A456B20F.png"
DAC_IMG = PKG / "references" / "FB27DA63CCC6BA9C819BC4ADEA17B9AC.png"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
    doc.add_paragraph()
    return table


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def set_doc_style(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.color.rgb = RGBColor(31, 78, 121)
        if name == "Heading 1":
            style.font.size = Pt(16)
        elif name == "Heading 2":
            style.font.size = Pt(13)
        else:
            style.font.size = Pt(11.5)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ETL4-7A35T XILINX ARTIX-7 FPGA 实验指导书")
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("实验九  ADC1采样与DAC1可调采样率回放实验")
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("默认4kSPS采样请求，支持通过参数调整采样率，用示波器对比ADC输入与DAC回放波形。")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(11)

    doc.add_paragraph()
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["实验名称", "ADC1采样 DAC1可调采样率回放实验"],
            ["适用开发板", "ETL4-7A35T / XC7A35T-2FGG484I"],
            ["接口类型", "I2C ADC + I2C DAC，共用一条I2C总线"],
            ["默认参数", "I2C目标400kHz，ADC_SAMPLE_HZ默认4000"],
            ["推荐测试", "40Hz到100Hz正弦波，2.0Vpp到2.4Vpp，1.65V偏置"],
        ],
    )
    note = doc.add_paragraph()
    note.add_run("注意：").bold = True
    note.add_run("本实验文档配套可调采样率源码工程包。若需要bit/ltx，请在Vivado 2020.1中重新生成。")


def build_doc() -> None:
    doc = Document()
    set_doc_style(doc)
    add_title_page(doc)

    doc.add_section(WD_SECTION_START.NEW_PAGE)

    doc.add_heading("实验简介", level=1)
    doc.add_paragraph(
        "本实验使用FPGA通过I2C总线读取ADC081C021的ADC1采样数据，并将最新采样值写入"
        "DAC081C085的DAC1输出端。实验使用双通道示波器同时观察信号源输入到ADC1的原始"
        "模拟波形，以及经过ADC采样、FPGA调度、DAC回放后的输出波形。"
    )
    doc.add_paragraph(
        "与固定10Hz演示版本不同，本版本将ADC采样请求速率抽象为ADC_SAMPLE_HZ参数。"
        "默认值为4kSPS，信号源输出频率不再由RTL写死，用户可以根据波形观感和I2C总线"
        "余量调整输入频率。"
    )

    doc.add_heading("9.1 原理与硬件连接", level=2)
    doc.add_paragraph(
        "ADC1采用ADC081C021，I2C地址为0x55，模拟输入由板上SMA接口接入。DAC1采用"
        "DAC081C085，I2C地址为0x0A，模拟输出由DAC1对应SMA接口输出。FPGA使用同一组"
        "I2C_SCL/I2C_SDA信号调度ADC读事务和DAC写事务。"
    )
    if ADC_IMG.exists():
        doc.add_picture(str(ADC_IMG), width=Cm(15.5))
        cap = doc.add_paragraph("图 9-1 ADC1/ADC2原理图与I2C地址")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if DAC_IMG.exists():
        doc.add_picture(str(DAC_IMG), width=Cm(15.5))
        cap = doc.add_paragraph("图 9-2 DAC1/DAC2原理图与I2C地址")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_table(
        doc,
        ["信号", "管脚", "电平标准", "说明"],
        [
            ["clk", "W19", "LVCMOS33", "50MHz系统时钟"],
            ["reset", "Y19", "LVCMOS33", "低电平复位输入"],
            ["i2c_scl", "W21", "LVCMOS33", "I2C时钟，开漏方式"],
            ["i2c_sda", "AA18", "LVCMOS33", "I2C数据，双向开漏方式"],
        ],
    )

    doc.add_heading("9.2 实验要求", level=2)
    for item in [
        "理解ADC输入频率、ADC采样请求速率和I2C总线带宽之间的区别。",
        "掌握通过ADC_SAMPLE_HZ参数调整采样率的方法。",
        "能够使用双通道示波器比较ADC1输入波形和DAC1回放波形。",
        "能够通过ILA观察dbg_adc_valid、dbg_adc_sample、dbg_dac_code、dbg_ack_error和dbg_replay_backlog_count。",
        "能够根据波形台阶、杂波、削峰现象判断频率和幅度设置是否合适。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("9.3 推荐参数与频率范围", level=2)
    doc.add_paragraph(
        "默认工程参数ADC_SAMPLE_HZ为4000，即4kSPS采样请求。示波器波形是否好看，"
        "主要取决于每个输入周期中有多少个采样点。一般建议每周期不少于40点作为演示波形，"
        "每周期20点左右可以观察趋势，少于10点时台阶会非常明显。"
    )
    add_table(
        doc,
        ["ADC_SAMPLE_HZ", "建议输入频率", "最高尝试频率", "说明"],
        [
            ["2000", "10Hz到50Hz", "100Hz", "最稳，等效上一版稳定演示思路"],
            ["4000", "10Hz到100Hz", "200Hz", "默认推荐，40Hz非常漂亮，100Hz可用但台阶较明显"],
            ["5000", "10Hz到125Hz", "250Hz", "接近400kHz共享I2C上限，必须看backlog"],
        ],
    )
    add_table(
        doc,
        ["信号源参数", "推荐值", "原因"],
        [
            ["波形", "正弦波优先，也可用三角波", "正弦波视觉上更平滑，便于观察台阶"],
            ["频率", "40Hz或50Hz用于展示，100Hz用于验证上限", "40Hz实测波形漂亮，100Hz已可用但台阶明显"],
            ["幅度", "2.0Vpp到2.4Vpp", "避免贴近0V和3.3V电源轨"],
            ["偏置", "1.65V", "让波形居中，减少上下削峰"],
        ],
    )

    doc.add_heading("9.4 实验步骤", level=2)
    for step in [
        "解压本实验压缩包，确认目录中包含rtl、xdc、build_project.tcl、run_synth.tcl、run_bitstream.tcl和docs目录。",
        "打开Vivado 2020.1，在Tcl Console中进入解压后的工程目录。",
        "执行source build_project.tcl新建工程，工程名为dac_adc1_loopback_variable_rate。",
        "如需调整采样率，打开rtl/dac_adc1_loopback_top.v，修改参数ADC_SAMPLE_HZ。",
        "执行source run_synth.tcl进行综合，确认无Error。",
        "执行source run_bitstream.tcl进行实现并生成bitstream，同时自动连接MARK_DEBUG信号到ILA。",
        "将信号源连接到ADC1输入SMA接口，将示波器CH1连接ADC1输入，CH2连接DAC1输出。",
        "信号源先设置为40Hz sine、2.4Vpp、offset 1.65V，下载bitstream后观察两路波形。",
        "将信号源改为100Hz sine、2.0Vpp、offset 1.65V，观察DAC输出是否仍可跟随输入波形。",
        "打开ILA，重点观察dbg_ack_error、dbg_replay_backlog和dbg_replay_backlog_count。",
    ]:
        add_number(doc, step)

    doc.add_heading("9.5 程序讲解", level=2)
    doc.add_paragraph(
        "顶层模块dac_adc1_loopback_top负责I2C开漏接口、I2C字节主机和ADC/DAC调度控制器例化。"
        "本版本增加了CLK_HZ、I2C_HZ、ADC_SAMPLE_HZ、DAC_ADDR和ADC_ADDR参数。"
    )
    doc.add_paragraph(
        "ADC_SAMPLE_HZ在顶层被转换为ADC_SAMPLE_DIV，计算方式为向上取整，确保采样请求接近"
        "目标频率。默认50MHz时钟、ADC_SAMPLE_HZ=4000时，ADC_SAMPLE_DIV为12500。"
    )
    doc.add_paragraph(
        "i2c_byte_master负责START、WRITE、READ_ACK、READ_NACK和STOP等字节级操作。"
        "I2C分频同样采用向上取整，避免整数截断导致实际SCL频率超过目标值。"
    )
    doc.add_paragraph(
        "dac_adc1_loopback_ctrl按ADC_SAMPLE_DIV产生ADC读取请求，读取ADC高低字节后取8位有效数据，"
        "再把最新采样值写入DAC1。若采样请求过快或DAC回放积压，dbg_replay_backlog和"
        "dbg_replay_backlog_count会提示总线调度已经跟不上。"
    )
    add_table(
        doc,
        ["调试信号", "含义", "正常现象"],
        [
            ["dbg_adc_sample", "最新ADC采样值", "随输入波形连续变化"],
            ["dbg_adc_valid", "ADC采样完成脉冲", "按采样率周期性跳变"],
            ["dbg_dac_code", "写入DAC1的回放码值", "跟随dbg_adc_sample变化"],
            ["dbg_ack_error", "I2C NACK粘滞标志", "应保持0"],
            ["dbg_replay_backlog_count", "采样/回放积压计数", "稳定参数下应保持0"],
        ],
    )

    doc.add_heading("9.6 实验现象", level=2)
    for item in [
        "40Hz正弦输入时，DAC1输出波形应与ADC1输入波形趋势一致，示波器显示较平滑，适合演示。",
        "100Hz正弦输入时，DAC1输出已可以跟随输入，但由于每周期采样点约40个，台阶会比40Hz明显。",
        "当输入频率继续提高时，波形台阶会增大；若采样点过少，示波器上可能看起来像阶梯或杂波。",
        "正常情况下dbg_ack_error保持0，dbg_replay_backlog_count保持0。",
        "若上下削峰，优先降低信号源Vpp，推荐从2.0Vpp、offset 1.65V重新测试。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("9.7 常见问题与思考题", level=2)
    for item in [
        "为什么40Hz比100Hz更漂亮？因为默认4kSPS下，40Hz约100点/周期，100Hz约40点/周期。",
        "为什么不能用3.3Vpp满幅输入？因为ADC/DAC工作在3.3V供电附近，满幅输入容易触碰上下电源轨，造成削峰。",
        "为什么不能直接把I2C_HZ改成3400000？因为真正的I2C HS-mode需要主机码进入流程、上拉和总线电容验证、SCL时序复核。",
        "如果dbg_replay_backlog_count增加，应该怎么处理？降低ADC_SAMPLE_HZ或优化I2C总线调度。",
        "如果100Hz台阶明显但功能正常，如何改善？降低输入频率到40Hz到50Hz，或提高采样率并确认backlog仍为0。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("工程文件说明", level=2)
    add_table(
        doc,
        ["目录或文件", "说明"],
        [
            ["00_老板先看这个.md", "快速测试方法和推荐参数"],
            ["rtl/", "Verilog源文件"],
            ["xdc/", "管脚约束文件"],
            ["build_project.tcl", "创建Vivado工程"],
            ["run_synth.tcl", "运行综合"],
            ["run_bitstream.tcl", "运行实现并生成bitstream"],
            ["setup_ila_current_project.tcl", "自动连接MARK_DEBUG信号到ILA"],
            ["docs/RATE_CONFIG.md", "采样率配置和频率范围说明"],
            ["docs/I2C_ADC采样率理论分析_20260602.md", "I2C带宽与理论采样率分析"],
        ],
    )

    OUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    PROJECT_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOC)
    doc.save(PROJECT_OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUT_DOC)
    print(PROJECT_OUTPUT)
