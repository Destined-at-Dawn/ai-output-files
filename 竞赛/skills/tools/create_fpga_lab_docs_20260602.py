from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE = Path("D:/AMD")
ADC_PROJECT = BASE / "dac_adc1_loopback_10hz"
EEPROM_PROJECT = Path("D:/xillin")
DOC_OUTPUT = BASE / "project_output" / "实验文档"

ADC_DOC = ADC_PROJECT / "实验指导书_ADC1采样DAC1回放固定10Hz.docx"
EEPROM_DOC = EEPROM_PROJECT / "实验指导书_AT21CS01单总线EEPROM读写暖复位.docx"

ADC_IMAGES = [
    BASE / "ADC" / "0178A2918E7353AF817F6876A456B20F.png",
    BASE / "ADC" / "FB27DA63CCC6BA9C819BC4ADEA17B9AC.png",
]
EEPROM_IMAGES = [
    BASE / "ADC" / "3C3A72DA4461A004095EE239D0A91E30.png",
]


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(10.5)
    set_run_font(r, "宋体")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, font_name="宋体", size=None, bold=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph, font_name="宋体", size=10.5):
    for run in paragraph.runs:
        set_run_font(run, font_name, size)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("第 ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    paragraph.add_run(" 页")


def set_document_style(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("上海锆赛电子有限公司")
    set_run_font(run, "宋体", 9)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        set_run_font(run, "宋体", 9)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, "黑体", 16 if level == 1 else 14 if level == 2 else 12, True)
    return p


def add_para(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, "宋体", 10.5, True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, "宋体", 10.5)
    else:
        r = p.add_run(text)
        set_run_font(r, "宋体", 10.5)
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.35
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, "宋体", 10.5)
        p.paragraph_format.line_spacing = 1.25


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r, "宋体", 10.5)
        p.paragraph_format.line_spacing = 1.25


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
    doc.add_paragraph()
    return table


def add_note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("注意：" + text)
    set_run_font(r, "宋体", 10.5, True)
    r.font.color.rgb = RGBColor(192, 0, 0)
    p.paragraph_format.first_line_indent = Pt(21)


def add_image(doc, image_path, caption, width_cm=15.5):
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Cm(width_cm))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        set_run_font(r, "宋体", 9)


def add_cover(doc, title, subtitle, project_name):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run("ETL4-7A35T XILINX ARTIX-7 FPGA 实验指导书")
    set_run_font(r, "黑体", 18, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(title)
    set_run_font(r, "黑体", 20, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run_font(r, "宋体", 12)

    doc.add_paragraph()
    add_table(doc, ["项目", "内容"], [
        ["适用平台", "XC7A35T-2FGG484I FPGA 学习板"],
        ["工程名称", project_name],
        ["软件环境", "Vivado 2020.1"],
        ["文档版本", "V1.0"],
        ["生成日期", "2026-06-02"],
    ])
    add_note(doc, "本文档配套工程已包含 RTL、XDC、Tcl、bit、ltx 文件，实验时请优先使用压缩包内工程。")
    doc.add_section(WD_SECTION.NEW_PAGE)


def add_common_vivado_steps(doc, xpr_path, bit_path, ltx_path):
    add_heading(doc, "实验步骤", 2)
    add_numbered(doc, [
        "解压本实验压缩包，确认工程目录中包含 rtl、xdc、vivado_project、build_project.tcl、run_synth.tcl、run_bitstream.tcl 等文件。",
        f"打开 Vivado 2020.1，选择 Open Project，打开工程文件：{xpr_path}。",
        "在 Flow Navigator 中依次检查 RTL Sources、Constraints 和顶层模块名称，确认没有缺失文件。",
        "如需重新生成工程，可在 Tcl Console 中进入工程目录后执行 source build_project.tcl。",
        "如需重新综合实现，可执行 source run_synth.tcl 与 source run_bitstream.tcl。",
        f"打开 Hardware Manager，连接开发板后选择 Program Device，bit 文件选择：{bit_path}。",
        f"需要观察 ILA 时，在同一对话框中加载 probes 文件：{ltx_path}。",
    ])


def add_project_files(doc, rows):
    add_heading(doc, "工程文件说明", 2)
    add_table(doc, ["文件或目录", "作用"], rows)


def create_adc_doc(path):
    doc = Document()
    set_document_style(doc)
    add_cover(
        doc,
        "实验八  ADC1 采样与 DAC1 固定 10Hz 回放实验",
        "使用 ADC1 采集外部 10Hz 模拟波形，并通过 DAC1 回放输出，在示波器上对比两路波形。",
        "dac_adc1_loopback_10hz",
    )

    add_heading(doc, "实验简介", 1)
    add_para(doc, "本实验使用 I2C 总线同时驱动 ADC081C021 与 DAC081C085。信号源向 ADC1 输入 10Hz 模拟波形，FPGA 以稳定的 2kHz 请求速率读取 ADC1 数据，并把最新采样值写入 DAC1。实验现象通过双通道示波器和 ILA 共同确认。")
    add_para(doc, "该版本的目标不是追求最高采样率，而是提供稳定、可复现、便于客户验证的 ADC 到 DAC 链路演示。")

    add_heading(doc, "8.1 原理与硬件连接", 2)
    add_para(doc, "ADC1 使用 ADC081C021，I2C 地址为 0x55，模拟输入从板上 SMA 接口接入。DAC1 使用 DAC081C085，I2C 地址为 0x0A，模拟输出从 DAC1 对应 SMA 接口输出。")
    for img, cap in zip(ADC_IMAGES, ["图 8-1 ADC1/ADC2 原理图与 I2C 地址", "图 8-2 DAC1/DAC2 原理图与 I2C 地址"]):
        add_image(doc, img, cap)
    add_table(doc, ["信号", "FPGA 管脚", "电平标准", "说明"], [
        ["clk", "W19", "LVCMOS33", "50MHz 系统时钟"],
        ["reset", "Y19", "LVCMOS33", "低电平复位"],
        ["i2c_scl", "W21", "LVCMOS33", "I2C 时钟线，开漏上拉"],
        ["i2c_sda", "AA18", "LVCMOS33", "I2C 数据线，开漏上拉"],
    ])
    add_table(doc, ["器件", "地址", "实验用途"], [
        ["ADC1 ADC081C021", "0x55", "采集外部 10Hz 模拟输入"],
        ["DAC1 DAC081C085", "0x0A", "回放 ADC1 最新采样值"],
        ["I2C 总线", "400kHz", "共享 ADC 读与 DAC 写事务"],
    ])

    add_heading(doc, "8.2 实验要求", 2)
    add_bullets(doc, [
        "掌握 I2C ADC 与 I2C DAC 在同一总线上的基本调度方式。",
        "能够使用 Vivado 下载 bit/ltx 文件并观察 ILA 采样信号。",
        "能够使用示波器双通道比较 ADC1 原始输入与 DAC1 回放输出。",
        "理解 2kHz 采样请求速率对 10Hz 波形约 200 点/周期的稳定演示意义。",
    ])
    add_common_vivado_steps(
        doc,
        "D:/AMD/dac_adc1_loopback_10hz/vivado_project/dac_adc1_loopback_10hz.xpr",
        "D:/AMD/dac_adc1_loopback_10hz/vivado_project/dac_adc1_loopback_10hz.runs/impl_1/dac_adc1_loopback_top.bit",
        "D:/AMD/dac_adc1_loopback_10hz/vivado_project/dac_adc1_loopback_10hz.runs/impl_1/dac_adc1_loopback_top.ltx",
    )
    add_heading(doc, "板上验证步骤", 3)
    add_numbered(doc, [
        "信号源输出 10Hz 波形，建议幅度不要贴近 0V 或 3.3V 电源轨，优先使用约 1.65V 偏置的小幅度波形。",
        "将信号源连接到 ADC1 输入 SMA 接口。",
        "将示波器 CH1 接 ADC1 输入，CH2 接 DAC1 输出。",
        "下载 bit 文件后观察两路波形，DAC1 输出应跟随 ADC1 输入变化。",
        "打开 ILA，重点观察 dbg_adc_sample、dbg_adc_valid、dbg_dac_code、dbg_ack_error、dbg_replay_backlog。",
    ])

    add_heading(doc, "8.3 程序讲解", 2)
    add_para(doc, "顶层模块 dac_adc1_loopback_top 实现 I2C SCL/SDA 开漏输出、例化 i2c_byte_master，并将 ADC/DAC 调度控制器的关键状态接入 ILA。")
    add_para(doc, "i2c_byte_master 是字节级 I2C 主机，负责 START、WRITE、READ_ACK、READ_NACK、STOP 等命令。当前工程实例化时设置 I2C_HZ 为 400000。")
    add_para(doc, "dac_adc1_loopback_ctrl 负责按 ADC_SAMPLE_DIV 产生采样请求。工程顶层覆盖参数 ADC_SAMPLE_DIV=25000，在 50MHz 时钟下对应 2kHz 请求速率。")
    add_table(doc, ["关键参数", "当前值", "说明"], [
        ["I2C_HZ", "400000", "I2C 快速模式，兼顾稳定性与调试余量"],
        ["ADC_SAMPLE_DIV", "25000", "50MHz/25000=2kHz"],
        ["ADC_ADDR", "7'h55", "ADC1 地址"],
        ["DAC_ADDR", "7'h0a", "DAC1 地址"],
    ])

    add_heading(doc, "8.4 实验现象", 2)
    add_bullets(doc, [
        "示波器 CH1 为 ADC1 输入的 10Hz 原始模拟波形。",
        "示波器 CH2 为 DAC1 回放输出，波形应随 CH1 同步变化。",
        "ILA 中 dbg_adc_valid 应周期性跳变，dbg_adc_sample 与 dbg_dac_code 应随输入波形变化。",
        "正常情况下 dbg_ack_error 应保持为 0；若 replay_backlog 长时间为 1，说明总线调度已经跟不上当前请求速率。",
    ])

    add_heading(doc, "8.5 常见问题与思考题", 2)
    add_table(doc, ["现象", "可能原因", "处理方法"], [
        ["DAC1 无输出", "I2C 未 ACK、地址或连线错误", "检查 dbg_ack_error、SCL/SDA 管脚和 DAC1 地址 0x0A"],
        ["回放波形削顶", "输入幅度贴近 ADC/DAC 电源轨", "降低输入幅度或设置 1.65V 偏置"],
        ["ILA 看不到完整 10Hz 周期", "采样窗口太短", "改用示波器看低频完整波形，ILA 只看数据变化"],
        ["回放抖动明显", "采样请求过快或总线拥塞", "保持 2kHz 稳定版本，先确认 ACK 错误为 0"],
    ])
    add_bullets(doc, [
        "为什么 2kHz 采样请求对 10Hz 波形已经有约 200 点/周期？",
        "为什么 I2C 共享总线不能简单等同于 ADC 芯片标称最高采样率？",
        "如果要采集更高频率波形，需要从总线、ADC 转换时间、DAC 建立时间哪些方面修改？",
    ])

    add_project_files(doc, [
        ["rtl/i2c_byte_master.v", "I2C 字节级主机"],
        ["rtl/dac_adc1_loopback_ctrl.v", "ADC1 采样与 DAC1 回放调度控制"],
        ["rtl/dac_adc1_loopback_top.v", "顶层端口、I2C 开漏和 ILA 信号"],
        ["xdc/dac_adc1_loopback_top.xdc", "时钟、复位、I2C 管脚和 ILA 约束"],
        ["vivado_project/.../dac_adc1_loopback_top.bit", "可直接下载的 bit 文件"],
        ["vivado_project/.../dac_adc1_loopback_top.ltx", "ILA probes 文件"],
    ])
    doc.save(path)


def create_eeprom_doc(path):
    doc = Document()
    set_document_style(doc)
    add_cover(
        doc,
        "实验九  AT21CS01 单总线 EEPROM 读写校验实验",
        "使用 FPGA 驱动 AT21CS01 单总线 EEPROM，写入固定数据后回读比较，并用 D0/D1 指示结果。",
        "eeprom_rw",
    )

    add_heading(doc, "实验简介", 1)
    add_para(doc, "本实验通过 FPGA 的 DQ 单线端口驱动 AT21CS01 EEPROM。程序自动扫描器件地址，完成 reset/discovery、标准速模式切换、数据写入、读回比对和 LED 结果显示。")
    add_para(doc, "当前交付版本针对板上按键复位后的暖复位问题做了修正：DQ 高电平等待 50ms，总线 reset 低电平 600us，reset release 高电平 200us，使断电重新下载和按键复位都能重新执行测试。")

    add_heading(doc, "9.1 原理与硬件连接", 2)
    add_para(doc, "AT21CS01 采用单线串行接口，DQ 既作为数据线，也通过上拉电阻获得工作能量。板上 DQ 通过 4.7kΩ 电阻上拉到 3.3V，因此复位后需要保证足够的高电平充电时间。")
    add_image(doc, EEPROM_IMAGES[0], "图 9-1 AT21CS01 单总线 EEPROM 原理图", 12.5)
    add_table(doc, ["信号", "FPGA 管脚", "电平标准", "说明"], [
        ["clk", "W19", "LVCMOS33", "50MHz 系统时钟"],
        ["reset", "Y19", "LVCMOS33", "低电平复位"],
        ["DQ", "Y6", "LVCMOS33", "AT21CS01 单总线数据/供电线"],
        ["D0", "AB22", "LVCMOS33", "成功指示灯，PASS 后点亮"],
        ["D1", "AB21", "LVCMOS33", "失败指示灯，FAIL 后点亮"],
    ])
    add_table(doc, ["时序项目", "当前值", "作用"], [
        ["DQ 高电平等待", "50ms", "给单线供电节点预充电"],
        ["Reset low", "600us", "覆盖冷启动和标准速状态下的暖复位"],
        ["Reset release", "200us", "复位释放后等待总线恢复"],
        ["EEPROM 写周期等待", "5ms", "等待内部写入完成"],
    ])

    add_heading(doc, "9.2 实验要求", 2)
    add_bullets(doc, [
        "掌握 AT21CS01 单总线 EEPROM 的开漏 DQ 控制方式。",
        "理解 reset/discovery、地址扫描、标准速模式切换和写读回比对流程。",
        "能够通过 D0/D1 直接判断 EEPROM 读写校验结果。",
        "能够使用 ILA 的 last_fail_state 与 last_fail_addr 定位失败阶段。",
    ])
    add_common_vivado_steps(
        doc,
        "D:/xillin/vivado_project/eeprom_rw.xpr",
        "D:/xillin/vivado_project/eeprom_rw.runs/impl_1/eeprom_rw_top.bit",
        "D:/xillin/vivado_project/eeprom_rw.runs/impl_1/eeprom_rw_top.ltx",
    )
    add_heading(doc, "板上验证步骤", 3)
    add_numbered(doc, [
        "下载 bit/ltx 文件后观察 D0/D1，复位过程中两灯均应熄灭。",
        "测试完成后，若 D0 亮且 D1 灭，表示写入与回读比较成功。",
        "按下板上 reset 键后再次观察，修正版应重新执行测试并回到 D0 亮。",
        "若 D1 亮，打开 ILA 观察 dbg_eeprom_state、dbg_current_addr、dbg_last_fail_state、dbg_last_fail_addr、dbg_eeprom_read_data。",
    ])

    add_heading(doc, "9.3 程序讲解", 2)
    add_para(doc, "顶层模块 eeprom_rw_top 负责连接 clk、reset、DQ、D0、D1，并通过 IOBUF 实现 DQ 开漏控制。LED 输出采用 done 使能：D0=done & pass，D1=done & fail，避免复位期间误亮。")
    add_para(doc, "at21cs01_master 是 EEPROM 主控制器。状态机先进行 DQ 高电平等待，然后执行总线 reset、discovery 响应检查、地址扫描、标准速模式命令、写入固定测试数据和读回比对。")
    add_para(doc, "测试数据为 A5、5A、3C、C3，从 EEPROM 地址 0x00 开始写入。程序会扫描 AT21CS01 的 3 位地址，只有全部地址都失败时才置 FAIL。")
    add_table(doc, ["关键调试信号", "说明"], [
        ["dbg_eeprom_state", "主状态机当前状态"],
        ["dbg_current_addr", "正在尝试的 AT21CS01 地址"],
        ["dbg_found_addr", "最终通过的地址"],
        ["dbg_eeprom_read_data", "读回的 4 字节数据"],
        ["dbg_last_fail_state", "最近一次 ACK/discovery 失败所在阶段"],
        ["dbg_last_fail_addr", "最近一次失败时的地址"],
        ["dbg_use_std_speed", "当前是否进入标准速访问阶段"],
    ])

    add_heading(doc, "9.4 实验现象", 2)
    add_bullets(doc, [
        "复位按下或复位刚释放时，D0 与 D1 均不亮。",
        "正常读写校验成功后，D0 常亮，D1 熄灭，ILA 中 pass=1。",
        "断电重新下载后应能 PASS；按键复位后也应能重新 PASS。",
        "若出现 D1 常亮，说明所有地址尝试后仍未完成写读回比对，需要查看 ILA 的 last_fail 信号。",
    ])

    add_heading(doc, "9.5 常见问题与思考题", 2)
    add_table(doc, ["现象", "可能原因", "处理方法"], [
        ["断电下载 PASS，按 reset 后 FAIL", "DQ 供电恢复不足或标准速状态下 reset 低脉宽不足", "使用本暖复位修正版，确认 50ms/600us/200us 参数生效"],
        ["D1 亮且 current_addr 到 7", "所有地址均未 ACK 或读回不匹配", "查看 last_fail_state，确认失败在 discovery、标准速命令还是 EEPROM 访问"],
        ["D0/D1 复位时异常亮", "LED 未被 done 信号门控", "确认顶层 D0=done&pass，D1=done&fail"],
        ["ILA 工程实现失败", "MARK_DEBUG 信号过多或深度过大", "减少宽计数器 probe，保持 C_DATA_DEPTH=4096"],
    ])
    add_bullets(doc, [
        "为什么 AT21CS01 不能直接按普通 Dallas 1-Wire 方式写状态机？",
        "为什么单线供电器件需要关注 DQ 高电平保持时间？",
        "为什么一次冷启动 PASS 不能证明按键复位也可靠？",
    ])

    add_project_files(doc, [
        ["rtl/at21cs01_master.v", "AT21CS01 reset/discovery、地址扫描、写读回比对主状态机"],
        ["rtl/eeprom_rw_top.v", "顶层 DQ IOBUF、LED 指示和 ILA 信号"],
        ["xdc/eeprom_rw_top.xdc", "clk/reset/DQ/D0/D1 管脚和 ILA 约束"],
        ["setup_ila_current_project.tcl", "生成精简 ILA，深度 4096"],
        ["vivado_project/.../eeprom_rw_top.bit", "可直接下载的 bit 文件"],
        ["vivado_project/.../eeprom_rw_top.ltx", "ILA probes 文件"],
        ["AT21CS01_protocol_fix_notes.md", "协议修正与暖复位修正记录"],
    ])
    doc.save(path)


def main():
    DOC_OUTPUT.mkdir(parents=True, exist_ok=True)
    create_adc_doc(ADC_DOC)
    create_eeprom_doc(EEPROM_DOC)
    # Keep central copies for the document archive.
    import shutil

    shutil.copy2(ADC_DOC, DOC_OUTPUT / ADC_DOC.name)
    shutil.copy2(EEPROM_DOC, DOC_OUTPUT / EEPROM_DOC.name)
    print(ADC_DOC)
    print(EEPROM_DOC)


if __name__ == "__main__":
    main()
