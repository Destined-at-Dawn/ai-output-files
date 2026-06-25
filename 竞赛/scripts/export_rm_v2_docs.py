from __future__ import annotations

import argparse
import re
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, Preformatted, SimpleDocTemplate, Spacer


DOCS = [
    ("00_v2_overview", "v2 总览", "README.md"),
    ("01_research_report", "深度调研与迭代报告", "research_report.md"),
    ("02_pack_readme", "学习包说明", "improved_pack/README.md"),
    ("03_technical_report", "技术文档", "improved_pack/docs/technical_report.md"),
    ("04_design", "设计思路", "improved_pack/docs/design.md"),
    ("05_code_walkthrough", "逐函数讲解", "improved_pack/docs/code_walkthrough.md"),
    ("06_step_by_step_tutorial", "一步步教程", "improved_pack/docs/tutorial_step_by_step.md"),
    ("07_test_plan", "测试计划", "improved_pack/docs/test_plan.md"),
    ("08_reflection_template", "创作心得模板", "improved_pack/docs/reflection.md"),
    ("09_submit_checklist", "提交检查清单", "improved_pack/docs/submit_checklist.md"),
]


def set_run_font(run, east_asia="微软雅黑", ascii_font="Arial", size=None, bold=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_doc_defaults(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)

    for name, size in [("Heading 1", 18), ("Heading 2", 15), ("Heading 3", 13)]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True

    for section in doc.sections:
        section.top_margin = Pt(56)
        section.bottom_margin = Pt(56)
        section.left_margin = Pt(62)
        section.right_margin = Pt(62)


def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = 1
    run = paragraph.add_run("第 ")
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)
    paragraph.add_run(" 页")


def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "")
    return text


def iter_md_blocks(text: str):
    in_code = False
    code_lines = []
    for raw in text.splitlines():
        line = raw.rstrip()
        if line.strip().startswith("```"):
            if in_code:
                yield ("code", "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        yield ("line", line)
    if code_lines:
        yield ("code", "\n".join(code_lines))


def add_docx_content(doc: Document, title: str, markdown: str, page_break_before=False):
    if page_break_before:
        doc.add_page_break()
    h = doc.add_heading(title, level=1)
    for run in h.runs:
        set_run_font(run, size=18, bold=True)

    for kind, value in iter_md_blocks(markdown):
        if kind == "code":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(value)
            set_run_font(run, east_asia="等线", ascii_font="Consolas", size=9)
            run.font.color.rgb = RGBColor(40, 40, 40)
            continue

        line = value.strip()
        if not line:
            continue
        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 3)
            text = clean_inline(line.lstrip("#").strip())
            p = doc.add_heading(text, level=level)
            for run in p.runs:
                set_run_font(run, size={1: 18, 2: 15, 3: 13}.get(level, 12), bold=True)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(clean_inline(line[2:].strip()))
            set_run_font(run)
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(clean_inline(re.sub(r"^\d+\.\s+", "", line)))
            set_run_font(run)
        elif line.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run = p.add_run(clean_inline(line.lstrip(">").strip()))
            set_run_font(run)
            run.italic = True
            run.font.color.rgb = RGBColor(90, 90, 90)
        elif line.startswith("|"):
            p = doc.add_paragraph()
            run = p.add_run(clean_inline(line))
            set_run_font(run, east_asia="等线", ascii_font="Consolas", size=9)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(18)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(clean_inline(line))
            set_run_font(run)


def make_docx(root: Path, docs: list[tuple[str, str, Path]], out_path: Path, combined=False):
    doc = Document()
    set_doc_defaults(doc)
    add_page_number(doc.sections[0])

    if combined:
        for idx, (_, title, path) in enumerate(docs):
            add_docx_content(doc, title, path.read_text(encoding="utf-8"), page_break_before=idx > 0)
    else:
        _, title, path = docs[0]
        add_docx_content(doc, title, path.read_text(encoding="utf-8"))

    doc.save(out_path)


def register_pdf_font() -> str:
    candidates = [
        Path(r"C:/Windows/Fonts/msyh.ttc"),
        Path(r"C:/Windows/Fonts/simhei.ttf"),
        Path(r"C:/Windows/Fonts/simsun.ttc"),
    ]
    for font in candidates:
        if font.exists():
            pdfmetrics.registerFont(TTFont("CNFont", str(font)))
            return "CNFont"
    return "Helvetica"


def pdf_styles(font_name: str):
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("title", parent=styles["Title"], fontName=font_name, fontSize=20, leading=26, spaceAfter=12),
        "h1": ParagraphStyle("h1", parent=styles["Heading1"], fontName=font_name, fontSize=17, leading=22, spaceBefore=10, spaceAfter=8),
        "h2": ParagraphStyle("h2", parent=styles["Heading2"], fontName=font_name, fontSize=14, leading=19, spaceBefore=8, spaceAfter=6),
        "h3": ParagraphStyle("h3", parent=styles["Heading3"], fontName=font_name, fontSize=12, leading=17, spaceBefore=6, spaceAfter=5),
        "body": ParagraphStyle("body", parent=styles["BodyText"], fontName=font_name, fontSize=10, leading=16, firstLineIndent=18, spaceAfter=4, alignment=TA_LEFT),
        "bullet": ParagraphStyle("bullet", parent=styles["BodyText"], fontName=font_name, fontSize=10, leading=15, leftIndent=18, bulletIndent=8, spaceAfter=3),
        "quote": ParagraphStyle("quote", parent=styles["BodyText"], fontName=font_name, fontSize=10, leading=15, leftIndent=18, textColor=colors.HexColor("#555555"), spaceAfter=5),
        "code": ParagraphStyle("code", parent=styles["Code"], fontName=font_name, fontSize=8.5, leading=11, leftIndent=8, backColor=colors.HexColor("#F4F4F4"), borderPadding=5, spaceBefore=4, spaceAfter=6),
    }


def add_pdf_content(story, title: str, markdown: str, styles, page_break_before=False):
    if page_break_before:
        story.append(PageBreak())
    story.append(Paragraph(escape(title), styles["title"]))
    story.append(Spacer(1, 3 * mm))

    for kind, value in iter_md_blocks(markdown):
        if kind == "code":
            story.append(Preformatted(value[:8000], styles["code"], maxLineLength=95))
            continue
        line = value.strip()
        if not line:
            continue
        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 3)
            text = escape(clean_inline(line.lstrip("#").strip()))
            story.append(Paragraph(text, styles[{1: "h1", 2: "h2", 3: "h3"}[level]]))
        elif line.startswith("- "):
            story.append(Paragraph("• " + escape(clean_inline(line[2:].strip())), styles["bullet"]))
        elif re.match(r"^\d+\.\s+", line):
            story.append(Paragraph(escape(clean_inline(line)), styles["bullet"]))
        elif line.startswith(">"):
            story.append(Paragraph(escape(clean_inline(line.lstrip(">").strip())), styles["quote"]))
        elif line.startswith("|"):
            story.append(Paragraph(escape(clean_inline(line)), styles["code"]))
        else:
            story.append(Paragraph(escape(clean_inline(line)), styles["body"]))


def make_pdf(root: Path, docs: list[tuple[str, str, Path]], out_path: Path, combined=False):
    font_name = register_pdf_font()
    styles = pdf_styles(font_name)
    story = []
    if combined:
        for idx, (_, title, path) in enumerate(docs):
            add_pdf_content(story, title, path.read_text(encoding="utf-8"), styles, page_break_before=idx > 0)
    else:
        _, title, path = docs[0]
        add_pdf_content(story, title, path.read_text(encoding="utf-8"), styles)

    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=out_path.stem,
    )
    doc.build(story)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("root", type=Path)
    args = parser.parse_args()

    root = args.root.resolve()
    docx_dir = root / "docs_docx"
    pdf_dir = root / "docs_pdf"
    docx_dir.mkdir(exist_ok=True)
    pdf_dir.mkdir(exist_ok=True)

    docs = [(file_stem, title, root / rel) for file_stem, title, rel in DOCS]
    missing = [str(path) for _, _, path in docs if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing markdown files: " + ", ".join(missing))

    for file_stem, title, path in docs:
        make_docx(root, [(file_stem, title, path)], docx_dir / f"{file_stem}.docx")
        make_pdf(root, [(file_stem, title, path)], pdf_dir / f"{file_stem}.pdf")

    make_docx(root, docs, docx_dir / "10_full_combined_document.docx", combined=True)
    make_pdf(root, docs, pdf_dir / "10_full_combined_document.pdf", combined=True)

    print(f"generated_docx={len(list(docx_dir.glob('*.docx')))}")
    print(f"generated_pdf={len(list(pdf_dir.glob('*.pdf')))}")


if __name__ == "__main__":
    main()
